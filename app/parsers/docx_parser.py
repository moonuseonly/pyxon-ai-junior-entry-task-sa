from pathlib import Path

from docx import Document as DocxDocument

from app.parsers.base import ParsedBlock, ParsedDocument


def _heading_level(style_name: str) -> int:
    # python-docx exposes Word's built-in styles as "Heading 1", "Heading 2", ...
    if style_name.startswith("Heading"):
        digits = "".join(ch for ch in style_name if ch.isdigit())
        return int(digits) if digits else 1
    if style_name == "Title":
        return 1
    return 0


class DOCXParser:
    def parse(self, file_path: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        blocks: list[ParsedBlock] = []
        full_text_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style and para.style.name else ""
            level = _heading_level(style_name)
            block_type = "heading" if level > 0 else "paragraph"
            blocks.append(ParsedBlock(text=text, block_type=block_type, level=level))
            full_text_parts.append(text)

        # Tables carry structured, often numeric, content — kept as distinct
        # blocks rather than flattened into prose so the analyzer/chunker can
        # treat them differently (e.g. never split a row across chunks).
        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            table_text = "\n".join(rows_text)
            if table_text.strip():
                blocks.append(ParsedBlock(text=table_text, block_type="table"))
                full_text_parts.append(table_text)

        return ParsedDocument(
            filename=Path(file_path).name,
            file_type="docx",
            blocks=blocks,
            full_text="\n".join(full_text_parts),
        )
    